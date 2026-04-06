import streamlit as st
import pypdf
from docx import Document
import json
import re
import time
import pandas as pd
from datetime import datetime
from groq import Groq
from openai import OpenAI

# Configuración de la página
st.set_page_config(page_title="Validador de Pagos UNAL - Centro de Prototipado", layout="wide")

st.markdown("""
    <style>
    .main, .stApp { background-color: #000c1a !important; color: #e6f1ff !important; }
    [data-testid="stSidebar"] { background-color: #001529 !important; border-right: 1px solid #1890ff; }
    .card { background-color: #001a33; padding: 20px; border-radius: 12px; border: 1px solid #004b8d; margin-bottom: 15px; }
    label, p, span, .stMarkdown, div[data-baseweb="select"] { color: #e6f1ff !important; }
    .stButton>button {
        background: linear-gradient(90deg, #1890ff 0%, #0050b3 100%) !important;
        color: white !important; border: none !important; border-radius: 10px !important;
        padding: 0.7rem 2.5rem !important; font-weight: bold !important;
    }
    h1, h2, h3 { color: #1890ff !important; font-family: 'Inter', sans-serif; }
    </style>
    """, unsafe_allow_html=True)

# --- UTILERÍAS ---
def normalize_money(val):
    if not val: return 0
    cleaned = re.sub(r'[\$\.\s,]', '', str(val))
    digits = re.sub(r'\D', '', cleaned)
    return int(digits) if digits else 0

def normalize_date(date_str):
    if not date_str: return "-"
    s = str(date_str).strip()
    
    # Manejo de fechas en español: "25 de marzo de 2026" → "2026/03/25"
    MESES_ES = {
        "enero":"01","febrero":"02","marzo":"03","abril":"04",
        "mayo":"05","junio":"06","julio":"07","agosto":"08",
        "septiembre":"09","octubre":"10","noviembre":"11","diciembre":"12"
    }
    m = re.search(r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', s.lower())
    if m:
        dia, mes_txt, anio = m.group(1), m.group(2), m.group(3)
        mes_num = MESES_ES.get(mes_txt.lower())
        if mes_num:
            return f"{anio}/{mes_num.zfill(2)}/{dia.zfill(2)}"

    # Formatos numéricos estándar
    clean = re.sub(r'[^\d/\-]', '', s)
    for fmt in ("%d/%m/%Y", "%Y/%m/%d", "%Y-%m-%d", "%d-%m-%Y"):
        try: return datetime.strptime(clean, fmt).strftime("%Y/%m/%d")
        except: pass
    return s

def normalize_period(period_str):
    """Normaliza periodos como 'marzo/2026', 'febrero/2026', '2026-03' a formato YYYY-MM"""
    if not period_str: return ""
    s = str(period_str).lower().strip()
    meses = {"enero":"01","febrero":"02","marzo":"03","abril":"04","mayo":"05","junio":"06",
             "julio":"07","agosto":"08","septiembre":"09","octubre":"10","noviembre":"11","diciembre":"12"}
    for mes, num in meses.items():
        if mes in s:
            year = re.search(r'\d{4}', s)
            return f"{year.group()}-{num}" if year else s
    # Ya puede estar en YYYY-MM
    m = re.match(r'(\d{4})-(\d{2})', s)
    if m: return f"{m.group(1)}-{m.group(2)}"
    return s

def parse_json_safe(content):
    """Limpia markdown y parsea JSON con tolerancia."""
    if not content: return {}
    if "```json" in content: content = content.split("```json")[1].split("```")[0]
    elif "```" in content: content = content.split("```")[1].split("```")[0]
    try: return json.loads(content.strip())
    except: return {}

# --- EXTRACCIÓN DE TEXTO ---
def extract_text_from_pdf(pdf_file):
    try:
        reader = pypdf.PdfReader(pdf_file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    except: return ""

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        from zipfile import ZipFile
        from io import BytesIO
        docx_file.seek(0)
        with ZipFile(BytesIO(docx_file.read())) as zf:
            if 'word/footer1.xml' in zf.namelist():
                footer_xml = zf.read('word/footer1.xml').decode('utf-8')
                import re
                footer_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', footer_xml)
                parts.extend([t.strip() for t in footer_texts if t.strip()])
            if 'word/header1.xml' in zf.namelist():
                header_xml = zf.read('word/header1.xml').decode('utf-8')
                header_texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', header_xml)
                parts.extend([t.strip() for t in header_texts if t.strip()])
        return "\n".join(parts)
    except Exception as e:
        return ""

# --- LÓGICA DE IA ESPECIALIZADA POR DOCUMENTO ---
def build_prompt(text, doc_type):
    if doc_type == "contrato":
        return f"""
    Para que el **Súper-Auditor de 120B** funcione correctamente, asegúrate de seguir estas reglas:
    
    1.  **📄 Contrato (PDF)**: Debe ser el archivo original del contrato u orden (OSE/CPS). De aquí la IA extraerá el número oficial y la vigencia.
    2.  **📋 Formato 4013 (PDF)**: 
        *   **Composición**: Este archivo es la unión (PDF) de:
            1. El excel `U-FT-12.010.069_Certificacion_determinacion_cedular_Rentas_de_Trabajo_V.6.1_VF` debidamente diligenciado.
            2. El comprobante de pago de Salud, Pensión y ARL.
            3. El certificado de la ARL.
        *   **Regla de Oro**: La unión de estos archivos debe llamarse exactamente **`4013AnexosOSE[Número].pdf`** (Ejemplo: `4013AnexosOSE14.pdf`).
        *   Asegúrate de que la tabla de fechas sea legible y contenga la columna de **'Pago'**.
    3.  **📝 Constancia (Docx)**: Debe ser el documento de cumplimiento en Word. El sistema verificará que el número de contrato y el periodo coincidan con los PDFs.
    
    *💡 El sistema realizará un **Triple Cruce** para asegurar que no haya discrepancias entre los tres archivos antes de que los radiques.*

CRITICO: Devuelve SOLO el bloque JSON. Sin texto adicional.
TEXTO DEL DOCUMENTO:
{text[:14000]}
"""
    elif doc_type == "4013":
        return f"""
Eres un auditor experto de la UNAL. Analiza este FORMATO 4013 PDF (puede tener varias páginas).

Extrae en JSON puro los siguientes campos:
- "numero_orden": Número de la orden contractual en la sección "2. RELACIÓN DE CONTRATOS". Extrae solo el número (ej: "14").
- "fecha_inicio_contratos": Fecha de inicio en sección "2. RELACIÓN DE CONTRATOS QUE ORIGINAN PAGOS (a)" (DD/MM/AAAA).
- "fecha_terminacion_contratos": Fecha de terminación en esa misma sección (DD/MM/AAAA).
- "periodo_solicitud_pago": Valor del campo "Periodo de solicitud de pago" (ej: "marzo/2026").
- "periodo_planilla": Valor del campo "Periodo de la planilla" (ej: "febrero/2026").
- "declaracion_formal_valor": El valor SI o NO que aparece en la celda de la DECLARACIÓN FORMAL PARA DISMINUCIÓN DE BASE DE RETENCIÓN.
- "total_aportes_obligatorios": CRITICO: Busca la sección "4. CALCULO APORTES OBLIGATORIOS AL SISTEMA DE SEGURIDAD SOCIAL INTEGRAL (SGSSI) MENSUALES". Dentro de esa tabla, encuentra la fila llamada "Total Aportes Obligatorios". Este valor es la SUMA de: IBC (Ingreso Base de Cotización) × 12.5% (Salud) + IBC × 16% (Pensión) + IBC × %ARL (según clase de riesgo). El resultado esperado es como 517300, 412800, 384000, etc. Extrae el valor de esa fila "Total Aportes Obligatorios" tal como aparece, solo el número (ejemplo: si ves "$517.300" extrae "517300", sin puntos ni signos).
- "total_pagado_planilla": Valor total pagado en la planilla de seguridad social (página 2). Solo el número total.
- "clave_planilla": En la página 2 (planilla de seguridad social) hay una tabla de liquidación con encabezado "Clave" que se divide en dos sub-columnas: "Pago" y "Planilla". El número bajo "Pago" tiene 8 dígitos (ej: 95049384). El número bajo "Planilla" tiene 10 dígitos (ej: 9500086013). DEBES EXTRAER SOLO EL VALOR DE LA COLUMNA "Planilla" (el de 10 dígitos).
- "fecha_pago_ss": ATENCION CRITICA: En la planilla hay una tabla con DOS fechas en columnas separadas: columna "Límite" y columna "Pago". El texto suele aparecer así: "Límite Pago\n06/03/2026 27/02/2026". El valor de "Límite" es el MES SIGUIENTE (la fecha de vencimiento), y el valor de "Pago" es la fecha REAL de pago. DEBES EXTRAER SOLO LA FECHA DE LA COLUMNA "PAGO" (la segunda fecha, la menor cronológicamente). NUNCA extraigas la fecha de la columna "Límite".
- "fecha_inicio_arl": Fecha inicio en el certificado ARL (al final del documento, tabla "Información general de la afiliación").
- "fecha_fin_arl": Fecha fin en el certificado ARL.
- "nombre_contratista": Nombre del contratista.
- "empresa_quipu": Número de empresa en QUIPU (ej: "4013").
- "riesgo_arl": Clase o nivel de riesgo ARL (solo el dígito, ej: "2").
- "fecha_diligenciamiento": Fecha en que fue diligenciado/firmado el formato 4013, que suele aparecer en la primera página como fecha de elaboración, firma o sello. Formato DD/MM/AAAA.

CRITICO: Devuelve SOLO el bloque JSON. Sin texto adicional.
TEXTO DEL DOCUMENTO:
{text[:16000]}
"""
    elif doc_type == "constancia":
        return f"""
Eres un auditor experto de la UNAL. Analiza esta CONSTANCIA DE CUMPLIMIENTO DOCX.

Extrae en JSON puro los siguientes campos:
- "numero_orden": El número de orden/contrato en el apartado "1. El contratista cumplió a satisfacción...". 
  En la tabla, busca la casilla "Número/Año" que tiene un valor como "14/2026" → extrae SOLO el primer número antes de "/" (ej: "14").
- "nombre_contratista": Nombre completo del contratista.
- "parcial_no": Número consecutivo del pago parcial actual (ej: 2).
- "informe_entregado": true si se menciona que se entregó informe, false si no.
- "fecha_pago_ss": Fecha de pago de seguridad social mencionada en el documento.
- "empresa_quipu": Número de empresa en QUIPU si aparece.
- "fecha_expedicion": CRITICO: Esta fecha aparece en el PIE del documento (al final), después de la frase 'Se expide la presente constancia para efectos del respectivo pago en la ciudad de Manizales, el día'. Por ejemplo: "Se expide la presente constancia para efectos del respectivo pago en la ciudad de Manizales, el día 25 de marzo de 2026" → extrae "25 de marzo de 2026". También puede aparecer como "el día 23 de marzo de 2026" o "Manizales, 23 de marzo de 2026". El texto puede estar separado por saltos de línea.
- "periodo_ss_constancia": En la sección "2. El contratista presentó la(s) planilla(s) número(s)..." hay una frase "para el (los) período(s) de 2026-02". Extrae SOLO el valor del período (ej: "2026-02" o "febrero/2026"). Este período debe coincidir con el "periodo_planilla" del formato 4013.

CRITICO: Devuelve SOLO el bloque JSON. Sin texto adicional.
TEXTO DEL DOCUMENTO:
{text[:15000]}
"""

def extract_data_with_ai(text, doc_type, model, api_key, provider):
    if not api_key: return {"error": "API Key faltante"}
    prompt = build_prompt(text, doc_type)
    try:
        if provider == "Groq":
            client = Groq(api_key=api_key)
        else:
            client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)

        for trial in range(2):
            try:
                params = {"model": model, "messages": [{"role": "user", "content": prompt}]}
                if provider == "Groq": params["response_format"] = {"type": "json_object"}
                response = client.chat.completions.create(**params)
                content = response.choices[0].message.content
                result = parse_json_safe(content)
                if result: return result
            except Exception as e:
                if trial == 0: time.sleep(4)
                else: return {"error": str(e)}
    except Exception as e:
        return {"error": str(e)}
    return {"error": "No se pudo extraer datos"}

# --- BARRA LATERAL ---
st.sidebar.title("🤖 Auditoría Multi-IA")
ai_provider = st.sidebar.selectbox("Proveedor", ["OpenRouter", "Groq"])

if ai_provider == "Groq":
    available_models = {"Súper-Auditor (GPT-OSS 120B)": "openai/gpt-oss-120b"}
    default_key = ""
else:
    tipo_filtro = st.sidebar.radio("Tipo de Modelos", ["Poderosos (SOTA)", "Versiones Gratis"])
    if tipo_filtro == "Poderosos (SOTA)":
        available_models = {
            "DeepSeek V3": "deepseek/deepseek-chat",
            "Gemini 2.0 Flash Lite (Paid)": "google/gemini-2.0-flash-lite-001"
        }
    else:
        available_models = {
            "GPT-OSS 120B (Free)": "openai/gpt-oss-120b:free",
            "Qwen 3.6 Plus": "qwen/qwen3.6-plus:free"
        }
    default_key = ""

selected_model_id = available_models[st.sidebar.selectbox("Modelo", list(available_models.keys()))]
user_api_key = st.sidebar.text_input(f"API Key {ai_provider}", value=default_key, type="password")

# --- UI PRINCIPAL ---
st.title("🛡️ Validador de Pagos - UNAL")
st.markdown("### Centro de Prototipado · Auditoría de Triple Cruce")

with st.expander("📖 Guía de Preparación de Documentos", expanded=False):
    st.info("""
    1. **Contrato (PDF)**: Archivo original de la orden OSE/CPS.
    2. **Formato 4013 (PDF)**: Unión de la certificación excel + comprobante planilla + certificado ARL.  
       Debe nombrarse: `4013AnexosOSE[N°].pdf` (Ej: `4013AnexosOSE14.pdf`)
    3. **Constancia (DOCX)**: Documento Word de cumplimiento contractual.
    """)

c1, c2, c3 = st.columns(3)
with c1: f_contrato = st.file_uploader("1. Contrato (PDF)", type=["pdf"])
with c2: f_4013 = st.file_uploader("2. Formato 4013 (PDF)", type=["pdf"])
with c3: f_constancia = st.file_uploader("3. Constancia (DOCX)", type=["docx"])

if st.button("🚀 Ejecutar Auditoría de Triple Cruce"):
    if not (f_contrato and f_4013 and f_constancia):
        st.error("⚠️ Debes cargar los 3 documentos.")
    else:
        with st.spinner(f"Analizando con {selected_model_id}... (3 análisis en secuencia)"):

            # 1. EXTRACCIÓN DE TEXTOS
            texts = {
                "contrato": extract_text_from_pdf(f_contrato),
                "4013": extract_text_from_pdf(f_4013),
                "constancia": extract_text_from_docx(f_constancia)
            }

            # 2. ANÁLISIS POR IA (uno por documento)
            extracted = {}
            progress = st.progress(0, text="Analizando Contrato...")
            extracted["contrato"] = extract_data_with_ai(texts["contrato"], "contrato", selected_model_id, user_api_key, ai_provider)
            progress.progress(33, text="Analizando Formato 4013...")
            extracted["4013"] = extract_data_with_ai(texts["4013"], "4013", selected_model_id, user_api_key, ai_provider)
            progress.progress(66, text="Analizando Constancia...")
            extracted["constancia"] = extract_data_with_ai(texts["constancia"], "constancia", selected_model_id, user_api_key, ai_provider)
            progress.progress(100, text="✅ Análisis completado")

            e_c = extracted.get("contrato", {})
            e_4 = extracted.get("4013", {})
            e_w = extracted.get("constancia", {})

            # 3. VALORES NORMALIZADOS
            num_contrato   = str(e_c.get("numero_orden") or "")
            num_4013       = str(e_4.get("numero_orden") or "")
            num_constancia = str(e_w.get("numero_orden") or "")

            # 4. VALIDACIÓN NOMBRE DE ARCHIVO (case-insensitive)
            st.divider()
            expected_fn = f"4013AnexosOSE{num_contrato}"
            actual_fn = f_4013.name.replace(".pdf", "").replace(" ", "")
            if actual_fn.lower() == expected_fn.lower():
                st.success(f"✅ Nombre del archivo 4013 correcto: `{f_4013.name}`")
            else:
                st.error(f"❌ Nombre incorrecto. Se esperaba: `{expected_fn}.pdf` · Se recibió: `{f_4013.name}`")

            # ================================================================
            # 5. FICHA ADMINISTRATIVA (TRIPLE CRUCE)
            # ================================================================
            st.divider()
            st.subheader("📋 Ficha Administrativa — Triple Cruce de Datos")

            def check(v1, v2, v3=None):
                vals = [str(v).strip().lower() for v in [v1, v2] if v and str(v) not in ["-", "None", "none", ""]]
                if v3: vals += [str(v3).strip().lower()]
                return len(set(vals)) <= 1

            def row_color(match):
                return ["background-color: #004d00; color: white"] * 4 if match else ["background-color: #800000; color: white"] * 4

            admin_rows = [
                ("N° Orden Contractual", num_contrato, num_4013, num_constancia),
                ("Nombre Contratista", e_c.get("nombre_contratista",""), e_4.get("nombre_contratista",""), e_w.get("nombre_contratista","")),
                ("Fecha Inicio Vigencia", normalize_date(e_c.get("fecha_inicio")), normalize_date(e_4.get("fecha_inicio_contratos")), "-"),
                ("Fecha Terminación", normalize_date(e_c.get("fecha_terminacion")), normalize_date(e_4.get("fecha_terminacion_contratos")), "-"),
                ("Empresa QUIPU", "4013", str(e_4.get("empresa_quipu","")), "-"),
            ]
            df_admin = pd.DataFrame(admin_rows, columns=["Campo", "Contrato PDF", "Formato 4013", "Constancia Word"]).astype(str)
            match_flags = [check(r[1], r[2], r[3]) for r in admin_rows]
            styled = df_admin.style.apply(lambda row: row_color(match_flags[row.name]), axis=1)
            st.table(styled)

            # ================================================================
            # 6. VALIDACIONES ESPECÍFICAS DEL FORMATO 4013
            # ================================================================
            st.divider()
            st.subheader("🔎 Validaciones Específicas del Formato 4013")
            v1, v2 = st.columns(2)

            # 6a. Periodo solicitud de pago vs Periodo planilla → debe coincidir con DECLARACIÓN FORMAL
            with v1:
                st.markdown("**Declaración Formal de Disminución de Base de Retención**")
                psp = normalize_period(e_4.get("periodo_solicitud_pago", ""))
                pp  = normalize_period(e_4.get("periodo_planilla", ""))
                decl = str(e_4.get("declaracion_formal_valor", "")).strip().upper()
                periodos_coinciden = (psp == pp) and bool(psp)
                esperado = "SI" if periodos_coinciden else "NO"

                col_a, col_b = st.columns(2)
                with col_a:
                    st.metric("Periodo solicitud pago", e_4.get("periodo_solicitud_pago", "N/A"))
                    st.metric("Periodo planilla",       e_4.get("periodo_planilla", "N/A"))
                with col_b:
                    st.metric("¿Coinciden?", "✅ SÍ" if periodos_coinciden else "❌ NO")
                    if decl == esperado:
                        st.success(f"Declaración Formal = **{decl}** ✅ Correcto")
                    else:
                        st.error(f"Declaración Formal = **{decl}** ❌ Debería ser **{esperado}**")

            # 6b. Fechas ARL vs Sección 2 del 4013
            with v2:
                st.markdown("**Fechas: Certificado ARL vs Relación de Contratos (Sección 2)**")
                fi_arl  = normalize_date(e_4.get("fecha_inicio_arl", ""))
                ff_arl  = normalize_date(e_4.get("fecha_fin_arl", ""))
                fi_s2   = normalize_date(e_4.get("fecha_inicio_contratos", ""))
                ff_s2   = normalize_date(e_4.get("fecha_terminacion_contratos", ""))

                col_c, col_d = st.columns(2)
                with col_c:
                    st.metric("F. Inicio ARL",    fi_arl)
                    st.metric("F. Inicio Sección 2", fi_s2)
                with col_d:
                    st.metric("F. Fin ARL",       ff_arl)
                    st.metric("F. Fin Sección 2", ff_s2)

                inicio_ok = (fi_arl == fi_s2) and fi_arl not in ["-", ""]
                fin_ok    = (ff_arl == ff_s2) and ff_arl not in ["-", ""]
                if inicio_ok: st.success("✅ Fechas de inicio coinciden")
                else:         st.error(f"❌ Inicio NO coincide: {fi_arl} ≠ {fi_s2}")
                if fin_ok:    st.success("✅ Fechas de terminación coinciden")
                else:         st.error(f"❌ Terminación NO coincide: {ff_arl} ≠ {ff_s2}")

            # ================================================================
            # 7. VALIDACIÓN APORTES OBLIGATORIOS vs PLANILLA
            # ================================================================
            st.divider()
            st.subheader("💰 Validación de Aportes Obligatorios (Sección 4 vs Planilla)")

            total_aportes  = normalize_money(e_4.get("total_aportes_obligatorios", 0))
            total_planilla = normalize_money(e_4.get("total_pagado_planilla", 0))

            # --- Fallback regex: buscar "Total Aportes Obligatorios" en el texto crudo del 4013 ---
            if total_aportes == 0:
                # Patrón 1: fila "Total Aportes Obligatorios" con valor en columna siguiente
                m_total = re.search(
                    r'Total\s+Aportes\s+Obligatorios\s*\$?\s*([\d.,]+)',
                    texts["4013"], re.IGNORECASE
                )
                if m_total:
                    total_aportes = normalize_money(m_total.group(1))
                
                # Patrón 2: buscar en tabla de sección 4 (valor grande de 6 dígitos)
                if total_aportes == 0:
                    m_alt = re.search(
                        r'(?:Total\s+Aportes|Total\s+Oblig)\s*(?:Obligatorios)?\s*[\n\r]+.*?([\d]{5,7})',
                        texts["4013"], re.IGNORECASE | re.DOTALL
                    )
                    if m_alt:
                        total_aportes = normalize_money(m_alt.group(1))

            # --- Fallback alternativo: sumar componentes del 4013 (Salud + Pensión + ARL) ---
            if total_aportes == 0:
                nums = []
                # Buscar valores de IBC y porcentajes en la sección 4
                ibc_match = re.search(r'IBC\s*\$?\s*([\d.,]+)', texts["4013"], re.IGNORECASE)
                if ibc_match:
                    ibc = normalize_money(ibc_match.group(1))
                    if 1000000 < ibc < 10000000:  # IBC válido entre 1M y 10M
                        # Calcular: 12.5% Salud + 16% Pensión + %ARL (típicamente 0.69% para riesgo 2)
                        salud_val = int(ibc * 0.125)
                        pension_val = int(ibc * 0.16)
                        arl_val = int(ibc * 0.0069)  # ~0.69% riesgo clase 2
                        total_aportes = salud_val + pension_val + arl_val
                        # Verificar si el resultado es razonable
                        if 200000 < total_aportes < 2000000:
                            nums = [salud_val, pension_val, arl_val]
                        else:
                            total_aportes = 0
                
                # Fallback: buscar valores individuales de salud, pension, arl en la tabla
                if total_aportes == 0:
                    for label in [r'salud\s*\$?\s*([\d.,]+)', r'pensi[oó]n\s*\$?\s*([\d.,]+)', r'ARL\s*\$?\s*([\d.,]+)', r'riesgos\s+laborales\s*\$?\s*([\d.,]+)']:
                        m = re.search(label, texts["4013"], re.IGNORECASE)
                        if m:
                            v = normalize_money(m.group(1))
                            if 50000 < v < 2000000:   # filtrar valores fuera de rango
                                nums.append(v)
                    if nums:
                        total_aportes = sum(nums)

            col_e, col_f, col_g = st.columns(3)
            with col_e: st.metric("Total Aportes Obligatorios (Sección 4)", f"${total_aportes:,}")
            with col_f: st.metric("Total Pagado en Planilla", f"${total_planilla:,}")
            with col_g:
                if total_aportes > 0 and total_planilla > 0:
                    if total_aportes <= total_planilla:
                        st.success(f"✅ Correcto: Aportes ≤ Planilla\n(${total_aportes:,} ≤ ${total_planilla:,})")
                    else:
                        st.error(f"❌ Error: Aportes > Planilla\n(${total_aportes:,} > ${total_planilla:,})")
                else:
                    st.warning("⚠️ No se pudieron extraer los valores para comparar")


            # ================================================================
            # 8. VALIDACIÓN PLANILLA PILA (en constancia) + FECHA PAGO
            # ================================================================
            st.divider()
            st.subheader("🧾 Validación de Planilla PILA y Fecha de Pago SS")
            v3a, v3b = st.columns(2)

            with v3a:
                st.markdown("**Planilla PILA**")
                pila = re.sub(r'\D', '', str(e_4.get("clave_planilla") or ""))
                # Fallback: buscar en el texto del PDF el número de 10 dígitos bajo la columna Planilla
                if len(pila) < 10:
                    # Patrón: busca secuencias de 10 dígitos típicas de planilla (start with 9)
                    fb = re.findall(r'(?<![\d])(9\d{9})(?![\d])', texts["4013"].replace(" ", "").replace("\n", " "))
                    if fb:
                        # Preferir el de 10 dígitos (Planilla), no el de 8 (Pago)
                        for candidate in fb:
                            if len(candidate) == 10:
                                pila = candidate
                                break
                encontrada = pila and pila in texts["constancia"].replace(" ", "")
                if encontrada: st.success(f"✅ Planilla `{pila}` encontrada en constancia")
                else:          st.error(f"❌ Planilla `{pila}` NO encontrada en constancia")

            with v3b:
                st.markdown("**Fechas de Pago SS**")
                f4_pago = normalize_date(e_4.get("fecha_pago_ss",""))
                fw_pago = normalize_date(e_w.get("fecha_pago_ss",""))
                st.write(f"Fecha SS en 4013: **{f4_pago}**")
                st.write(f"Fecha SS en Constancia: **{fw_pago}**")
                if f4_pago not in ["-",""] and f4_pago == fw_pago: st.success("✅ Fechas coinciden")
                elif f4_pago in ["-",""] or fw_pago in ["-",""]:   st.warning("⚠️ No se encontraron fechas en ambos documentos")
                else: st.error(f"❌ Fechas no coinciden: {f4_pago} ≠ {fw_pago}")

            # ================================================================
            # 9. VALIDACIÓN FECHA EXPEDICIÓN CONSTANCIA vs DILIGENCIAMIENTO 4013 y FECHA PAGO SS
            # ================================================================
            st.divider()
            st.subheader("📅 Validación de Fechas de Expedición y Diligenciamiento")

            fecha_expedicion       = normalize_date(e_w.get("fecha_expedicion", ""))
            fecha_diligenciamiento = normalize_date(e_4.get("fecha_diligenciamiento", ""))
            fecha_pago_planilla    = normalize_date(e_4.get("fecha_pago_ss", ""))

            # Fallback robusto: extraer fecha_expedicion directamente del texto si la IA falló
            if fecha_expedicion in ["-", ""]:
                constancia_text = texts.get("constancia", "")
                patrones = [
                    r'el\s+d[i\xed]a\s+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|\d{1,2}[/\-]\d{2}[/\-]\d{4})',
                    r'ciudad\s+de\s+Manizales[,\s]+?(el\s+)?(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})',
                    r'Manizales[,\s]+?(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})',
                    r'expide[^\d]+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})',
                    r'se\s+expide[^\d]+(\d{1,2}[/\-]\d{2}[/\-]\d{4})',
                    r'día\s+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})',
                    r'(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})',
                    r'(\d{1,2}[/\-]\d{2}[/\-]\d{4})',
                ]
                for patron in patrones:
                    m_exp = re.search(patron, constancia_text, re.IGNORECASE | re.DOTALL)
                    if m_exp:
                        fecha_expedicion = normalize_date(m_exp.group(-1).strip())
                        break

            col_h, col_i, col_j = st.columns(3)
            with col_h:
                st.metric("Fecha Expedición Constancia", fecha_expedicion or "No encontrada")
                st.caption("'Se expide la presente constancia...'")
            with col_i:
                st.metric("Fecha Diligenciamiento 4013", fecha_diligenciamiento or "No encontrada")
                st.caption("Fecha de elaboración/firma del 4013")
            with col_j:
                st.metric("Fecha Pago Planilla SS", fecha_pago_planilla or "No encontrada")
                st.caption("Fecha real de pago en planilla PILA")

            # Verificación 1: Expedición constancia == Diligenciamiento 4013
            if fecha_expedicion not in ["-", ""] and fecha_diligenciamiento not in ["-", ""]:
                if fecha_expedicion == fecha_diligenciamiento:
                    st.success(f"✅ Fecha expedición constancia ({fecha_expedicion}) coincide con fecha de diligenciamiento del 4013")
                else:
                    st.error(f"❌ Las fechas NO coinciden: Constancia={fecha_expedicion} · 4013={fecha_diligenciamiento}")
            else:
                st.warning("⚠️ No se encontró alguna de las fechas (expedición constancia o diligenciamiento 4013)")

            # Verificación 2: Expedición constancia > Fecha pago planilla SS
            if fecha_expedicion not in ["-", ""] and fecha_pago_planilla not in ["-", ""]:
                try:
                    dt_exp  = datetime.strptime(fecha_expedicion, "%Y/%m/%d")
                    dt_pago = datetime.strptime(fecha_pago_planilla, "%Y/%m/%d")
                    if dt_exp > dt_pago:
                        st.success(f"✅ Fecha expedición ({fecha_expedicion}) es POSTERIOR a la fecha de pago SS ({fecha_pago_planilla})")
                    else:
                        st.error(f"❌ La fecha de expedición ({fecha_expedicion}) NO es posterior a la fecha de pago SS ({fecha_pago_planilla})")
                except:
                    st.warning("⚠️ No se pudieron comparar las fechas (formato desconocido)")
            else:
                st.warning("⚠️ No se encontró la fecha de expedición o la fecha de pago de planilla para comparar")

            # ================================================================
            # 10. VALIDACIÓN PERÍODO SS CONSTANCIA vs PERÍODO PLANILLA 4013
            # ================================================================
            st.divider()
            st.subheader("📆 Validación de Período de Seguridad Social (Constancia vs 4013)")

            periodo_ss_constancia = normalize_period(str(e_w.get("periodo_ss_constancia", "")))
            periodo_planilla_4013 = normalize_period(str(e_4.get("periodo_planilla", "")))

            # Fallback regex: buscar "para el (los) período(s) de YYYY-MM" en la constancia
            if periodo_ss_constancia in ["", "none", "none"]:
                m_periodo = re.search(
                    r'per[ií]odo\(s\)\s+de\s+(\d{4}-\d{2})',
                    texts["constancia"],
                    re.IGNORECASE
                )
                if m_periodo:
                    periodo_ss_constancia = normalize_period(m_periodo.group(1))

            col_k, col_l, col_m = st.columns(3)
            with col_k:
                st.metric("Período SS en Constancia", periodo_ss_constancia or "No encontrado")
                st.caption("Sección 2: 'para el (los) período(s) de...'")
            with col_l:
                st.metric("Período Planilla 4013", periodo_planilla_4013 or "No encontrado")
                st.caption("Campo 'Periodo de la planilla'")
            with col_m:
                if periodo_ss_constancia and periodo_planilla_4013:
                    if periodo_ss_constancia == periodo_planilla_4013:
                        st.success(f"✅ Períodos coinciden ({periodo_ss_constancia})")
                    else:
                        st.error(f"❌ Períodos NO coinciden: Constancia={periodo_ss_constancia} · 4013={periodo_planilla_4013}")
                else:
                    st.warning("⚠️ No se encontraron los períodos para comparar")

            # ================================================================
            # JSON DEBUG
            # ================================================================
            with st.expander("🔍 Ver datos JSON extraídos por la IA"):
                st.json(extracted)
